from typing import Any, Dict, Optional
from pathlib import Path
import tempfile

try:
    from docx import Document
    from docx.shared import Pt, Cm
except ImportError:  # pragma: no cover - docx may not be installed
    Document = None
    Pt = None

class LegalDocumentGenerator:
    """Generate legal filings in Traditional Chinese."""

    def __init__(self, case_info: Dict[str, Any]):
        self.case_info = case_info
        if Document is None:
            raise RuntimeError(
                "python-docx is required to generate documents. Please install it via 'pip install python-docx'."
            )
        self.doc = Document()
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "標楷體"
        font.size = Pt(16)
        style.paragraph_format.line_spacing = Pt(24)  # 1.5 line spacing

        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def build(self) -> None:
        self.doc.add_heading(self.case_info.get('title', '起訴狀'), level=1)
        self._add_basic_info()
        self._add_claims()
        self._add_facts()
        self._add_laws()
        self._add_evidence()
        self._add_attachments()

    def save(self, filepath: str) -> None:
        self.doc.save(filepath)

    def save_pdf(self, filepath: str) -> None:
        """Save the generated document as a PDF using docx2pdf if available."""
        try:
            from docx2pdf import convert
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise RuntimeError(
                "docx2pdf is required for PDF export. Install it via 'pip install docx2pdf'."
            ) from exc

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "temp.docx"
            self.doc.save(docx_path)
            convert(str(docx_path), filepath)

    # Internal helpers
    def _add_basic_info(self) -> None:
        self.doc.add_paragraph(f"案號：{self.case_info.get('case_number', '')}")
        self.doc.add_paragraph(f"當事人：{self.case_info.get('parties', '')}")
        self.doc.add_paragraph(f"法院：{self.case_info.get('court', '')}")

    def _add_claims(self) -> None:
        self.doc.add_paragraph("壹、訴之聲明")
        self.doc.add_paragraph(self.case_info.get('claims', ''))

    def _add_facts(self) -> None:
        self.doc.add_paragraph("貳、事實與理由")
        self.doc.add_paragraph(self.case_info.get('facts', ''))

    def _add_laws(self) -> None:
        self.doc.add_paragraph("參、法律依據")
        for law in self.case_info.get('laws', []):
            self.doc.add_paragraph(f"• {law}")

    def _add_evidence(self) -> None:
        self.doc.add_paragraph("肆、證據目錄")
        for ev in self.case_info.get('evidence', []):
            self.doc.add_paragraph(f"【{ev['id']}】{ev['summary']}")

    def _add_attachments(self) -> None:
        attachments = self.case_info.get("attachments")
        if not attachments:
            return
        self.doc.add_paragraph("伍、附件")
        for att in attachments:
            desc = att.get("description", "")
            att_id = att.get("id", "")
            self.doc.add_paragraph(f"【{att_id}】{desc}")


def create_filing(
    case_info: Dict[str, Any],
    output_path: str,
    pdf_path: Optional[str] = None,
) -> None:
    """Helper function to quickly generate a filing document.

    Parameters
    ----------
    case_info : Dict[str, Any]
        Information about the case. The dictionary may include an
        ``attachments`` key for an optional **附件** section.
    output_path : str
        Location to write the DOCX file.
    pdf_path : Optional[str]
        If provided, also export the document to this PDF path.

    The generated document uses 標楷體 font, 2.5 cm margins, and 1.5 line
    spacing.
    """
    generator = LegalDocumentGenerator(case_info)
    generator.build()
    generator.save(output_path)
    if pdf_path:
        generator.save_pdf(pdf_path)

